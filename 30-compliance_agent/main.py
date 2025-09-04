import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional
from agno.agent import Agent
from agno.models.openai import OpenAIChat
from agno.tools.exa import ExaTools
from agno.tools.reasoning import ReasoningTools
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv

class AnalyseurConformiteUE:
    """Système d'Analyse de Conformité UE"""
    
    def __init__(self, config_path: Optional[str] = None):
        """Initialise l'analyseur de conformité"""
        load_dotenv(config_path)
        
        # Récupération des clés API
        self.exa_api_key = os.getenv("EXA_API_KEY")
        self.openai_api_key = os.getenv("OPENAI_API_KEY")
        
        # Validation des clés
        if not self.openai_api_key:
            raise ValueError("Clé API OpenAI manquante")
        
        # Initialisation de l'agent
        self._configurer_agent()
    
    def _configurer_agent(self):
        """Configure l'agent d'analyse de conformité"""
        outils = [ReasoningTools(add_instructions=True)]
        
        # Ajout d'ExaTools si la clé est disponible
        if self.exa_api_key:
            outils.append(ExaTools(
                api_key=self.exa_api_key,
                include_domains=["eur-lex.europa.eu", "ec.europa.eu"]
            ))
        
        self.agent = Agent(
            name="Analyseur de Conformité UE",
            model=OpenAIChat(
                id="gpt-4-turbo", 
                api_key=self.openai_api_key
            ),
            tools=outils,
            show_tool_calls=True,
            description="""
            Agent spécialisé dans l'analyse de la conformité réglementaire européenne.
            Fournit des évaluations de conformité structurées pour les entreprises
            opérant sur le marché européen.
            """,
            instructions=self._obtenir_instructions_agent(),
            markdown=True,
            stream=True
        )
    
    def _obtenir_instructions_agent(self) -> str:
        """Instructions détaillées pour l'agent de conformité"""
        return """
        Vous êtes un expert en conformité UE. Votre mission est de fournir des analyses
        réglementaires complètes pour les entreprises opérant sur les marchés européens.

        CADRE D'ANALYSE :
        1. PROFIL DE L'ENTREPRISE
           - Identifier les activités principales
           - Déterminer la classification sectorielle
           - Évaluer l'étendue géographique des opérations UE
           
        2. CARTOGRAPHIE RÉGLEMENTAIRE
           - Rechercher les directives et règlements applicables
           - Identifier les exigences sectorielles spécifiques
           - Suivre les calendriers de mise en œuvre
           
        3. ÉVALUATION DES RISQUES DE CONFORMITÉ
           - Catégoriser les règlements par criticité (Élevée/Moyenne/Faible)
           - Identifier les lacunes potentielles de conformité
           - Évaluer les sanctions et pénalités
           
        4. DOMAINES RÉGLEMENTAIRES PRINCIPAUX :
           - Protection des données : RGPD, ePrivacy
           - IA & Numérique : Loi sur l'IA, Loi sur les services numériques
           - Environnement : Pacte vert, CSRD, Taxonomie
           - Sectoriel : Automobile, Services financiers
           - Sécurité des produits : Marquage CE, Cybersécurité
           
        5. STRUCTURE DU RAPPORT :
           - Résumé exécutif avec priorités clés
           - Calendrier réglementaire
           - Matrice des risques
           - Feuille de route de conformité
           - Exigences en ressources
           
        STRATÉGIE DE RECHERCHE :
        - Utiliser uniquement des sources officielles UE
        - Inclure des références spécifiques aux articles
        - Fournir les dates de publication et d'entrée en vigueur
        - Vérifier la cohérence entre les sources
        """
    
    def analyser_entreprise(self, nom_entreprise: str, secteur: str, 
                          contexte_supplementaire: str = "") -> str:
        """Analyse les exigences de conformité pour une entreprise spécifique"""
        
        requete = f"""
        Effectuez une analyse complète de conformité réglementaire UE pour :
        
        Entreprise : {nom_entreprise}
        Secteur : {secteur}
        Contexte supplémentaire : {contexte_supplementaire}
        
        Identifiez tous les règlements UE applicables, leurs exigences de conformité,
        les délais de mise en œuvre et les impacts potentiels sur l'activité.
        Fournissez des conseils de conformité spécifiques et actionnables.
        """
        
        print(f"🔍 Début de l'analyse de conformité pour {nom_entreprise}...")
        response = self.agent.print_response(requete)
        
        return response
    
    def generer_rapport_word(self, chemin_markdown: str, 
                           chemin_sortie: Optional[str] = None) -> str:
        """Convertit un rapport markdown en document Word formaté"""
        
        # Lecture du contenu markdown
        try:
            with open(chemin_markdown, "r", encoding="utf-8") as fichier:
                contenu = fichier.read()
        except FileNotFoundError:
            raise FileNotFoundError(f"Fichier markdown introuvable : {chemin_markdown}")
        
        # Initialisation du document Word
        doc = Document()
        
        # Ajout de la page de titre
        self._ajouter_page_titre(doc)
        
        # Traitement du contenu
        self._traiter_contenu_markdown(doc, contenu)
        
        # Ajout du pied de page
        self._ajouter_pied_document(doc)
        
        # Sauvegarde du document
        if not chemin_sortie:
            horodatage = datetime.now().strftime("%Y%m%d_%H%M%S")
            chemin_sortie = f"Rapport_Conformite_UE_{horodatage}.docx"
        
        doc.save(chemin_sortie)
        print(f"📄 Rapport Word sauvegardé : {chemin_sortie}")
        
        return chemin_sortie
    
    def _ajouter_page_titre(self, doc: Document):
        """Ajoute une page de titre professionnelle"""
        # Titre
        titre = doc.add_heading("Analyse de Conformité Réglementaire UE", 0)
        titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Sous-titre
        sous_titre = doc.add_paragraph("Rapport d'Intelligence Réglementaire")
        sous_titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sous_titre_format = sous_titre.runs[0].font
        sous_titre_format.size = Pt(14)
        sous_titre_format.italic = True
        
        # Date
        date_para = doc.add_paragraph(f"Généré le : {datetime.now().strftime('%d %B %Y')}")
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Saut de page
        doc.add_page_break()
    
    def _traiter_contenu_markdown(self, doc: Document, contenu: str):
        """Traite le contenu markdown et l'ajoute au document Word"""
        lignes = contenu.splitlines()
        
        for ligne in lignes:
            ligne = ligne.strip()
            if not ligne:
                continue
            
            # En-têtes
            if ligne.startswith("# "):
                doc.add_heading(ligne[2:], level=1)
            elif ligne.startswith("## "):
                doc.add_heading(ligne[3:], level=2)
            elif ligne.startswith("### "):
                doc.add_heading(ligne[4:], level=3)
            elif ligne.startswith("#### "):
                doc.add_heading(ligne[5:], level=4)
            
            # Listes
            elif ligne.startswith("* ") or ligne.startswith("- "):
                self._ajouter_puce(doc, ligne[2:].strip())
            elif re.match(r'^\d+\. ', ligne):
                self._ajouter_numero(doc, re.sub(r'^\d+\. ', '', ligne))
            
            # Paragraphes normaux
            else:
                self._ajouter_paragraphe_formate(doc, ligne)
    
    def _ajouter_puce(self, doc: Document, texte: str):
        """Ajoute une puce formatée"""
        p = doc.add_paragraph(style="List Bullet")
        self._ajouter_texte_formate(p, texte)
    
    def _ajouter_numero(self, doc: Document, texte: str):
        """Ajoute un élément numéroté"""
        p = doc.add_paragraph(style="List Number")
        self._ajouter_texte_formate(p, texte)
    
    def _ajouter_paragraphe_formate(self, doc: Document, texte: str):
        """Ajoute un paragraphe avec formatage"""
        p = doc.add_paragraph()
        self._ajouter_texte_formate(p, texte)
    
    def _ajouter_texte_formate(self, paragraphe, texte: str):
        """Ajoute du texte avec formatage gras/italique"""
        # Traitement du texte en gras (**texte**)
        parties = re.split(r"(\*\*.*?\*\*)", texte)
        for partie in parties:
            if partie.startswith("**") and partie.endswith("**"):
                run = paragraphe.add_run(partie[2:-2])
                run.bold = True
            else:
                # Traitement du texte en italique (*texte*)
                parties_italique = re.split(r"(\*.*?\*)", partie)
                for partie_italique in parties_italique:
                    if (partie_italique.startswith("*") and partie_italique.endswith("*") 
                        and not partie_italique.startswith("**")):
                        run = paragraphe.add_run(partie_italique[1:-1])
                        run.italic = True
                    else:
                        paragraphe.add_run(partie_italique)
    
    def _ajouter_pied_document(self, doc: Document):
        """Ajoute un pied de document avec métadonnées"""
        doc.add_page_break()
        doc.add_heading("Informations du Document", level=2)
        
        infos_pied = [
            f"Généré le : {datetime.now().strftime('%d-%m-%Y %H:%M:%S')}",
            "Source : Sources officielles UE (eur-lex.europa.eu, ec.europa.eu)",
            "Avertissement : Ce rapport est à des fins d'information uniquement.",
            "Recommandation : Consultez un conseil juridique pour des questions de conformité spécifiques."
        ]
        
        for info in infos_pied:
            p = doc.add_paragraph(info)
            p.style = "Caption"

def main():
    """Fonction d'exécution principale"""
    try:
        # Initialisation de l'analyseur
        analyseur = AnalyseurConformiteUE()
        
        # Exemple d'utilisation
        nom_entreprise = "Renault"
        secteur = "Industrie Automobile"
        contexte_supplementaire = "Focus sur les systèmes IA, réglementations d'émissions et reporting de durabilité"
        
        # Lancement de l'analyse
        print("🚀 Début de l'analyse de conformité UE...")
        reponse = analyseur.analyser_entreprise(
            nom_entreprise=nom_entreprise,
            secteur=secteur,
            contexte_supplementaire=contexte_supplementaire
        )
        
        # Conversion en Word (en supposant qu'un fichier markdown a été généré)
        chemin_markdown = "conformite_automobile.md"
        if Path(chemin_markdown).exists():
            chemin_rapport_word = analyseur.generer_rapport_word(chemin_markdown)
            print(f"✅ Analyse terminée ! Rapport sauvegardé : {chemin_rapport_word}")
        else:
            print("⚠️ Fichier markdown introuvable. Veuillez vérifier le chemin.")
            
    except Exception as e:
        print(f"❌ Erreur lors de l'analyse : {e}")
        raise

if __name__ == "__main__":
    main()